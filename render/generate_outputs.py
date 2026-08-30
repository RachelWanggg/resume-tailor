#!/usr/bin/env python3
"""
generate_outputs.py  —  1-page .docx + .pdf from a tailored resume .md

Every visual decision comes from resume.css, which is the single source of styling truth. The
.docx template is only a container: build() strips all of its content and re-applies margins and
fonts from the CSS, so a blank document works just as well as a hand-styled one.

Usage:  python3 render/generate_outputs.py <input.md> [--no-fit] [--scale=N]
                                           [--template FILE] [--outdir DIR]
Output: <stem>.docx and <stem>.pdf
        The input .md is deleted afterward only when it lives in the outputs directory.

PDF export needs either Microsoft Word (macOS) or LibreOffice (any platform).
"""

import re, sys, shutil, shutil as _shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

REPO_ROOT = Path(__file__).resolve().parent.parent
CSS_PATH = Path(__file__).with_name('resume.css')

# Optional .docx to inherit fonts/numbering from. Ships absent: python-docx's own default template
# is used instead, which keeps the repo self-contained. Override with --template to carry your own
# styling. Never point this outside the repo — that breaks the tool for everyone but you.
DEFAULT_TEMPLATE = Path(__file__).with_name('template.docx')
TEMPLATE = DEFAULT_TEMPLATE if DEFAULT_TEMPLATE.exists() else None


def _new_document(out_path):
    """Open the style template, or a blank document when no template is configured."""
    if TEMPLATE is not None:
        if not TEMPLATE.exists():
            raise SystemExit(f'Error: template not found: {TEMPLATE}')
        shutil.copy(TEMPLATE, out_path)
        return Document(out_path)
    return Document()

def _load_css_rules():
    """Load the small, intentional CSS contract used by this renderer."""
    css = CSS_PATH.read_text(encoding='utf-8')
    rules = {}
    for selector, body in re.findall(r'([^{}]+)\{([^{}]*)\}', css):
        props = {}
        for name, value in re.findall(r'([\w-]+)\s*:\s*([^;]+)', body):
            props[name.strip()] = value.strip()
        for part in selector.split(','):
            rules[part.strip()] = props
    return rules

CSS = _load_css_rules()

def _css(selector, prop, default=None):
    return CSS.get(selector, {}).get(prop, default)

def _css_pt(value, default):
    if not value:
        return default
    m = re.match(r'^([0-9.]+)\s*(pt|px)$', value)
    if not m:
        return default
    number, unit = float(m.group(1)), m.group(2)
    return number if unit == 'pt' else number * 0.75

def _css_hex(value, default):
    m = re.search(r'#[0-9a-fA-F]{6}', value or '')
    if m:
        return m.group(0)[1:].upper()
    return default

def _page_margins():
    value = _css('@page', 'margin', '0.6in')
    nums = [float(x) for x in re.findall(r'([0-9.]+)in', value)]
    if len(nums) == 1:
        return (nums[0],) * 4
    if len(nums) == 2:
        return nums[0], nums[1], nums[0], nums[1]
    if len(nums) == 4:
        return tuple(nums)
    return 0.6, 0.7, 0.6, 0.7

CSS_TOP, CSS_RIGHT, CSS_BOTTOM, CSS_LEFT = _page_margins()

# ── exact values from base_ai_engineer.docx inspection ───────────────────────
BLACK       = RGBColor(0x00, 0x00, 0x00)
LINK_CLR    = _css_hex(_css('a', 'color'), '0563C1')
BORDER_CLR  = _css_hex(_css('h2', 'border-bottom'), '1A5CA8')
FONT_MAIN   = _css('body', 'font-family', 'Calibri').split(',')[0].strip()
FONT_MAIN   = FONT_MAIN.strip('"\'')
NAME_PT     = _css_pt(_css('h1', 'font-size'), 16)
CONTACT_PT  = _css_pt(_css('h1 + p', 'font-size'), 8)
BODY_PT     = _css_pt(_css('body', 'font-size'), 9)
BULLET_PT   = _css_pt(_css('body', 'font-size'), 9)
SECTION_PT  = _css_pt(_css('h2', 'font-size'), 11)
RIGHT_TAB   = int((8.5 - CSS_LEFT - CSS_RIGHT) * 1440)
BUL_LEFT    = int(_css_pt(_css('ul', 'padding-left'), 13) * 20)  # CSS px → twips
BUL_HANG    = int(12 * 20)  # marker offset within the CSS list indent
NUM_ID      = '2'     # bullet list numId defined in base_ai_engineer.docx

# ── auto-fit scaling ─────────────────────────────────────────────────────────
# Every font size / spacing value above is a *base* value at SCALE = 1.0.
# main() searches for the largest SCALE that still fits on one page, so a short
# resume grows to fill the sheet instead of leaving the bottom third blank.
SCALE      = 1.0
BASE_SIZES = dict(NAME_PT=NAME_PT, CONTACT_PT=CONTACT_PT, BODY_PT=BODY_PT,
                  BULLET_PT=BULLET_PT, SECTION_PT=SECTION_PT,
                  BUL_LEFT=BUL_LEFT, BUL_HANG=BUL_HANG)

def _half(x):
    """Round to the nearest half-point — the finest size Word stores."""
    return round(x * 2) / 2

def apply_scale(s):
    """Rescales all type sizes / indents / line spacing to factor `s`."""
    global SCALE, NAME_PT, CONTACT_PT, BODY_PT, BULLET_PT, SECTION_PT
    global BUL_LEFT, BUL_HANG, LINE_TIGHT
    SCALE      = s
    NAME_PT    = _half(BASE_SIZES['NAME_PT']    * s)
    CONTACT_PT = _half(BASE_SIZES['CONTACT_PT'] * s)
    BODY_PT    = _half(BASE_SIZES['BODY_PT']    * s)
    BULLET_PT  = _half(BASE_SIZES['BULLET_PT']  * s)
    SECTION_PT = _half(BASE_SIZES['SECTION_PT'] * s)
    BUL_LEFT   = int(BASE_SIZES['BUL_LEFT'] * s)
    BUL_HANG   = int(BASE_SIZES['BUL_HANG'] * s)
    # Word line-spacing units are 240ths of a line; CSS line-height is relative.
    LINE_TIGHT = int(float(re.sub(r'[^0-9.]', '', _css('body', 'line-height', '1.18'))) * 240)

# ── inline helpers ────────────────────────────────────────────────────────────

# Tokenizes **bold** and [label](url) markdown inline, in document order.
INLINE_RE = re.compile(r'\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)')

def parse_inline(text):
    """Returns list of (text, bold, url) preserving order."""
    segs = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            segs.append((text[pos:m.start()], False, None))
        if m.group(1) is not None:
            segs.append((m.group(1), True, None))
        else:
            segs.append((m.group(2), False, m.group(3)))
        pos = m.end()
    if pos < len(text):
        segs.append((text[pos:], False, None))
    return segs

def strip_tags(t):
    return re.sub(r'<[^>]+>', '', t)

def strip_italic_marker(text):
    """Detects a whole-string *italic* wrapper (single asterisks) and returns (text, is_italic)."""
    m = re.match(r'^\*([^*].*[^*]|[^*])\*$', text.strip())
    if m:
        return m.group(1), True
    return text, False

def _sp(p, before=0, after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(_half(before * SCALE))
    pf.space_after  = Pt(_half(after  * SCALE))

LINE_TIGHT = int(float(re.sub(r'[^0-9.]', '', _css('body', 'line-height', '1.18'))) * 240)

def _line(p, val=None):
    val = LINE_TIGHT if val is None else val
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'), str(val)); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)

def _rtab(p):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(RIGHT_TAB))
    tabs.append(tab); pPr.append(tabs)

def _border_bottom(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), BORDER_CLR)
    pBdr.append(bot); pPr.append(pBdr)

def _runs(p, text, font=None, pt=None):
    """Renders **bold** and [label](url) inline markdown as runs/hyperlinks."""
    pt = BODY_PT if pt is None else pt
    font = FONT_MAIN if font is None else font
    for seg, bold, url in parse_inline(strip_tags(text)):
        if not seg:
            continue
        if url:
            _add_hyperlink(p, seg, url, bold=bold, pt=pt, font=font)
            continue
        r = p.add_run(seg)
        r.bold = bold
        r.font.name = font
        r.font.size = Pt(pt)
        r.font.color.rgb = BLACK

def _jc(p, val='both'):
    pPr = p._p.get_or_add_pPr()
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), val)
    pPr.append(jc)

# ── paragraph builders ────────────────────────────────────────────────────────

def add_name(doc, text):
    p = doc.add_paragraph()
    _sp(p, 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(strip_tags(text))
    r.bold = True; r.font.name = FONT_MAIN; r.font.size = Pt(NAME_PT); r.font.color.rgb = BLACK

def add_contact(doc, text):
    p = doc.add_paragraph()
    _sp(p, 0, 10); _border_bottom(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    parts = [s.strip() for s in strip_tags(text).split('|') if s.strip()]

    def _run(txt):
        r = p.add_run(txt); r.font.name = FONT_MAIN; r.font.size = Pt(CONTACT_PT); r.font.color.rgb = BLACK; return r

    for i, part in enumerate(parts):
        if i > 0:
            _run(' | ')
        is_email = '@' in part and ' ' not in part
        is_url   = part.startswith('http') or 'linkedin.com' in part or 'github.com' in part
        if is_email:
            _add_hyperlink(p, part, 'mailto:' + part, bold=False, pt=CONTACT_PT, font=FONT_MAIN)
        elif is_url:
            url = part if part.startswith('http') else 'https://' + part
            _add_hyperlink(p, part, url, bold=False, pt=CONTACT_PT, font=FONT_MAIN)
        else:
            _run(part)

def add_section(doc, text):
    p = doc.add_paragraph()
    _sp(p, 4, 1); _jc(p); _border_bottom(p)
    r = p.add_run(text)
    r.bold = True; r.font.name = FONT_MAIN; r.font.size = Pt(SECTION_PT); r.font.color.rgb = RGBColor.from_string(_css_hex(_css('h2', 'color'), '000000'))

def add_company(doc, left, right):
    p = doc.add_paragraph(); _sp(p, 2, 0); _jc(p); _rtab(p); _line(p)
    text = strip_tags(left)
    whole_wrap = re.match(r'^\*\*(.+)\*\*$', text)
    if whole_wrap:
        # whole line was one ** ** span (possibly with a nested [label](url)) — force bold on all pieces
        segs = [(s, True, u) for s, _, u in parse_inline(whole_wrap.group(1))]
    else:
        # mixed bold, e.g. "**Northlake University**, MS in Computer Science" — respect per-segment bold
        segs = parse_inline(text)
    for seg, bold, url in segs:
        if not seg: continue
        if url:
            _add_hyperlink(p, seg, url, bold=bold, pt=BODY_PT, font=FONT_MAIN)
            continue
        r = p.add_run(seg); r.bold = bold; r.font.name = FONT_MAIN; r.font.size = Pt(BODY_PT); r.font.color.rgb = BLACK
    if right:
        rt = p.add_run('\t'); rt.font.name = FONT_MAIN; rt.font.size = Pt(BODY_PT)
        dr = p.add_run(right); dr.bold = False; dr.font.name = FONT_MAIN; dr.font.size = Pt(BODY_PT); dr.font.color.rgb = BLACK

def add_role(doc, left, right):
    p = doc.add_paragraph(); _sp(p, 0, 0); _jc(p); _rtab(p); _line(p)
    text, italic = strip_italic_marker(strip_tags(left))
    rl = p.add_run(text); rl.font.name = FONT_MAIN; rl.font.size = Pt(BODY_PT); rl.font.color.rgb = BLACK
    if italic: rl.italic = True
    if right:
        rt = p.add_run('\t'); rt.font.name = FONT_MAIN; rt.font.size = Pt(BODY_PT)
        rr = p.add_run(right); rr.font.name = FONT_MAIN; rr.font.size = Pt(BODY_PT); rr.font.color.rgb = BLACK

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Paragraph')
    _sp(p, 0, 0); _line(p)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(BUL_LEFT)); ind.set(qn('w:hanging'), str(BUL_HANG))
    pPr.append(ind)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'both'); pPr.append(jc)
    numPr = OxmlElement('w:numPr')
    ilvl  = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), NUM_ID)
    numPr.append(ilvl); numPr.append(numId)
    pPr.insert(0, numPr)
    _runs(p, text, font=FONT_MAIN, pt=BULLET_PT)

def _add_hyperlink(p, text, url, bold=True, pt=None, font=None):
    pt = BODY_PT if pt is None else pt
    font = FONT_MAIN if font is None else font
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font); rFonts.set(qn('w:hAnsi'), font); rFonts.set(qn('w:cs'), font)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(pt * 2))); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(pt * 2))); rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement('w:b'))
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    color = OxmlElement('w:color'); color.set(qn('w:val'), LINK_CLR); rPr.append(color)
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; run.append(t)
    hl.append(run); p._p.append(hl)

def add_proj_title(doc, text, url=None, suffix=''):
    p = doc.add_paragraph(); _sp(p, 2, 0); _jc(p); _rtab(p)
    if url:
        _add_hyperlink(p, text, url, bold=True, pt=BODY_PT)
        if suffix:
            sp_run = p.add_run(' '); sp_run.font.name = FONT_MAIN
            _runs(p, suffix, pt=BODY_PT)
    else:
        r = p.add_run(strip_tags(text)); r.bold = True; r.font.name = FONT_MAIN; r.font.size = Pt(BODY_PT); r.font.color.rgb = BLACK

def add_skills(doc, text):
    p = doc.add_paragraph(); _sp(p, 0, 0); _jc(p)
    if ':' in text:
        label, rest = text.split(':', 1)
        rl = p.add_run(strip_tags(label) + ':'); rl.bold = True; rl.font.name = FONT_MAIN; rl.font.size = Pt(BODY_PT); rl.font.color.rgb = BLACK
        rr = p.add_run(strip_tags(rest)); rr.bold = False; rr.font.name = FONT_MAIN; rr.font.size = Pt(BODY_PT); rr.font.color.rgb = BLACK
    else:
        _runs(p, text, pt=BODY_PT)

# ── document builder ──────────────────────────────────────────────────────────

META_RE = re.compile(r'<span class="meta">(.*?)</span>')

# Reviewer diagnostics are useful to the orchestrator but must never appear in
# the submitted resume.  Keep this guard in the renderer as a final safety net
# in case a writer leaves diagnostics after the resume body.
INTERNAL_DIAGNOSTIC_PREFIXES = (
    'HARD_MISMATCH:', 'BULLET_BUDGET:', 'PRESET:', 'NEIGHBOR:',
    'DOMAIN_WORDS_LANDED:', 'STACK_LANDING:', 'SUBSTITUTIONS:',
    'SKILLS_ONLY:', 'PROJECTS_ONLY:', 'TRUE_GAPS:',
)

def build(md_text, out_path):
    doc = _new_document(out_path)
    body = doc.element.body

    # Remove orphaned header/footer references from sectPr
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        for tag in ('w:headerReference', 'w:footerReference'):
            for el in sectPr.findall(qn(tag)):
                sectPr.remove(el)

    # Clear all body paragraphs/tables (keep sectPr)
    for child in list(body):
        t = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if t in ('p', 'tbl', 'sdt'):
            body.remove(child)

    # Margins — exact values from base_ai_engineer.docx
    for sec in doc.sections:
        sec.top_margin    = Inches(CSS_TOP)
        sec.bottom_margin = Inches(CSS_BOTTOM)
        sec.left_margin   = Inches(CSS_LEFT)
        sec.right_margin  = Inches(CSS_RIGHT)

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('BANK_UPDATE:'): break
        if line == '---' and i+1 < len(lines) and lines[i+1].startswith('BANK_UPDATE:'): break
        if line.startswith(INTERNAL_DIAGNOSTIC_PREFIXES): break

        if line.startswith('# '):
            add_name(doc, line[2:].strip()); i += 1; continue

        # Contact line: email | phone | links | city  (also legacy [REDACTED])
        if line.startswith('[REDACTED]') or (
            '|' in line and ('@' in line or 'linkedin.com' in line or 'github.com' in line)):
            add_contact(doc, strip_tags(line)); i += 1; continue

        if line.startswith('## '):
            add_section(doc, line[3:].strip()); i += 1; continue

        if line.strip() in ('---', '') :
            i += 1; continue

        if line.startswith('- '):
            add_bullet(doc, line[2:].strip()); i += 1; continue

        # Standalone **[Text](url)** hyperlink project title (with optional trailing text)
        proj_link_m = re.match(r'^\*\*\[([^\]]+)\]\(([^)]+)\)\*\*(.*)', line.strip())
        if proj_link_m:
            add_proj_title(doc, proj_link_m.group(1), url=proj_link_m.group(2),
                           suffix=proj_link_m.group(3).strip()); i += 1; continue

        # Standalone **Bold** project title
        if re.match(r'^\*\*[^*]+\*\*$', line.strip()):
            add_proj_title(doc, re.sub(r'\*\*', '', line.strip())); i += 1; continue

        # Skills line **Label:** value
        if re.match(r'^\*\*[^*]+:\*\*', line.strip()):
            add_skills(doc, re.sub(r'\*\*', '', line.strip())); i += 1; continue

        # Paragraph with meta spans (company or role lines)
        metas = META_RE.findall(line)
        left  = META_RE.sub('', line).strip()
        if metas:
            if '**' in line or re.search(r'\[[^\]]+\]\([^)]+\)', line):
                add_company(doc, left, metas[0])
            else:
                add_role(doc, left, metas[0])
        else:
            p = doc.add_paragraph(); _sp(p, 0, 0)
            _runs(p, line.strip())
        i += 1

    doc.save(out_path)

# ── PDF export + page-fill measurement ───────────────────────────────────────

BOTTOM_MARGIN_PT = 0.6 * 72   # matches sec.bottom_margin in build()
TOP_MARGIN_PT    = 0.6 * 72

WORD_APP = Path('/Applications/Microsoft Word.app')


def _soffice_bin():
    """Locate a LibreOffice binary, including the macOS bundle path."""
    found = _shutil.which('soffice') or _shutil.which('libreoffice')
    if found:
        return found
    mac_bundle = Path('/Applications/LibreOffice.app/Contents/MacOS/soffice')
    return str(mac_bundle) if mac_bundle.exists() else None


def pdf_backend():
    """Which converter is available: 'word', 'libreoffice', or None."""
    if sys.platform == 'darwin' and WORD_APP.exists():
        return 'word'
    if _soffice_bin():
        return 'libreoffice'
    return None


def _to_pdf_word(docx_path, pdf_path):
    import subprocess
    script = f'''
    tell application "{WORD_APP}"
    open POSIX file "{docx_path}"
    set theDoc to active document
    save as theDoc file name (POSIX file "{pdf_path}" as string) file format format PDF
    close theDoc saving no
end tell
'''
    r = subprocess.run(['osascript', '-e', script], capture_output=True, text=True)
    return (r.returncode == 0 and pdf_path.exists()), r.stderr.strip()


def _to_pdf_libreoffice(docx_path, pdf_path):
    import subprocess
    soffice = _soffice_bin()
    r = subprocess.run(
        [soffice, '--headless', '--convert-to', 'pdf', '--outdir',
         str(pdf_path.parent), str(docx_path)],
        capture_output=True, text=True,
    )
    # LibreOffice names the output after the input stem; rename if that differs.
    produced = pdf_path.parent / f'{docx_path.stem}.pdf'
    if produced.exists() and produced != pdf_path:
        produced.replace(pdf_path)
    return (r.returncode == 0 and pdf_path.exists()), r.stderr.strip()


def to_pdf(docx_path, pdf_path):
    """Exports docx → pdf via Microsoft Word (macOS) or LibreOffice. Returns (ok, stderr)."""
    if pdf_path.exists():
        pdf_path.unlink()
    backend = pdf_backend()
    if backend == 'word':
        return _to_pdf_word(docx_path, pdf_path)
    if backend == 'libreoffice':
        return _to_pdf_libreoffice(docx_path, pdf_path)
    return False, (
        'No PDF converter found. Install LibreOffice (https://www.libreoffice.org/download/) '
        'or, on macOS, Microsoft Word. The .docx was still written and can be exported by hand.'
    )

def measure(pdf_path):
    """Returns (page_count, fill) where fill is the fraction of the printable
    column height on page 1 that actually holds content."""
    import pymupdf
    with pymupdf.open(pdf_path) as doc:
        pages = doc.page_count
        page  = doc[0]
        bottoms = [b[3] for b in page.get_text('blocks')]
        for d in page.get_drawings():
            bottoms.append(d['rect'].y1)
        if not bottoms:
            return pages, 0.0
        usable = page.rect.height - TOP_MARGIN_PT - BOTTOM_MARGIN_PT
        return pages, (max(bottoms) - TOP_MARGIN_PT) / usable

MAX_SCALE = 1.25   # body caps at ~11.25pt — bigger reads like a school assignment

def autofit(md, out_docx, out_pdf, target=0.97, max_iters=7):
    """Searches for the largest SCALE whose PDF is exactly one page and whose
    content reaches `target` of the printable height. Returns the chosen scale
    (None if Word export is unavailable and no measurement could be made)."""
    s, lo, hi, best = 1.0, None, None, None
    for it in range(max_iters):
        apply_scale(s)
        build(md, out_docx)
        ok, err = to_pdf(out_docx, out_pdf)
        if not ok:
            print(f'PDF generation failed: {err}')
            return best[0] if best else None
        pages, fill = measure(out_pdf)
        print(f'  fit pass {it+1}: scale={s:.3f} pages={pages} fill={fill:.3f}')
        if pages == 1:
            if best is None or s > best[0]:
                best = (s, fill)
            if fill >= target or s >= MAX_SCALE:
                return s
            lo = s
            nxt = min(MAX_SCALE, s * min(1.30, target / max(fill, 0.30)))
            if hi is not None:
                nxt = min(nxt, (s + hi) / 2)
        else:
            hi = s
            nxt = s * 0.93 if lo is None else (lo + s) / 2
        if abs(nxt - s) < 0.004:
            break
        s = nxt
    return best[0] if best else None

# ── main ──────────────────────────────────────────────────────────────────────

def main():
    global TEMPLATE
    args = sys.argv[1:]
    argv = [a for a in args if not a.startswith('--')]
    flags = {a for a in args if a.startswith('--')}

    # --template FILE / --template=FILE
    tmpl = next((f.split('=', 1)[1] for f in flags if f.startswith('--template=')), None)
    if tmpl is None and '--template' in args:
        idx = args.index('--template')
        if idx + 1 < len(args):
            tmpl = args[idx + 1]
            if tmpl in argv:
                argv.remove(tmpl)
    if tmpl:
        TEMPLATE = Path(tmpl).expanduser().resolve()

    # --outdir DIR / --outdir=DIR, defaulting to the repo's outputs/ directory
    odir = next((f.split('=', 1)[1] for f in flags if f.startswith('--outdir=')), None)
    if odir is None and '--outdir' in args:
        idx = args.index('--outdir')
        if idx + 1 < len(args):
            odir = args[idx + 1]
            if odir in argv:
                argv.remove(odir)

    if not argv:
        print('Usage: python3 render/generate_outputs.py <input.md> [--no-fit] [--scale=N] '
              '[--template FILE] [--outdir DIR]')
        raise SystemExit(1)
    md_path = Path(argv[0]).resolve()
    if not md_path.exists():
        print(f'Error: {md_path} not found'); raise SystemExit(1)

    outputs_dir = (Path(odir).expanduser().resolve() if odir
                   else (REPO_ROOT / 'outputs'))
    outputs_dir.mkdir(parents=True, exist_ok=True)

    md = md_path.read_text()
    md = re.split(r'\n---\nBANK_UPDATE:', md)[0]

    out     = outputs_dir / f'{md_path.stem}.docx'
    pdf_out = outputs_dir / f'{md_path.stem}.pdf'

    fixed = next((f for f in flags if f.startswith('--scale=')), None)
    if '--no-fit' in flags or fixed:
        apply_scale(float(fixed.split('=')[1]) if fixed else 1.0)
        build(md, out)
        ok, err = to_pdf(out, pdf_out)
        print(f'DOCX → {out}')
        print(f'PDF  → {pdf_out}' if ok else f'PDF generation failed: {err}')
    else:
        chosen = autofit(md, out, pdf_out)
        if chosen is None:
            # No usable measurement (Word unavailable) — fall back to base sizes.
            apply_scale(1.0); build(md, out)
            print(f'DOCX → {out}')
        else:
            if abs(chosen - SCALE) > 1e-6:   # last pass wasn't the winner — rebuild it
                apply_scale(chosen)
                build(md, out)
                ok, err = to_pdf(out, pdf_out)
                if not ok:
                    print(f'PDF generation failed: {err}')
            pages, fill = measure(pdf_out)
            print(f'DOCX → {out}')
            print(f'PDF  → {pdf_out}  (scale {chosen:.3f}, {fill*100:.0f}% of page filled)')

    # Only ever delete the source markdown when it is scratch input sitting in the outputs
    # directory. A file the user passed in from anywhere else is theirs, and stays.
    if md_path.parent == outputs_dir:
        md_path.unlink(); print(f'Removed {md_path.name}')

if __name__ == '__main__':
    main()
